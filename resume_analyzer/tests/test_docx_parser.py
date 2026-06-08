"""Tests for Word document parser module."""

from io import BytesIO

import pytest

from resume_analyzer.services.document.docx_parser import DocxParser
from resume_analyzer.utils.file_utils import is_resume_upload, resume_file_kind


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Build a minimal in-memory .docx with resume-like text."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Email: jane@example.com")
    doc.add_paragraph("Skills: Python, SQL, Docker")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestDocxParser:
    """Word parser unit tests."""

    def test_validate_invalid_bytes(self) -> None:
        """Corrupted bytes should fail validation."""
        parser = DocxParser()
        valid, _ = parser.validate_docx(b"not a docx")
        assert valid is False

    def test_extract_text_from_docx(self, sample_docx_bytes: bytes) -> None:
        """Parser extracts paragraph text from a valid docx."""
        parser = DocxParser()
        text = parser.extract_from_upload(sample_docx_bytes)
        assert "Jane Doe" in text
        assert "Python" in text


class TestResumeFileHelpers:
    """Resume upload extension helpers."""

    def test_is_resume_upload(self) -> None:
        """Supported extensions are recognized."""
        assert is_resume_upload("resume.pdf")
        assert is_resume_upload("resume.DOCX")
        assert is_resume_upload("resume.doc")
        assert not is_resume_upload("resume.txt")

    def test_resume_file_kind(self) -> None:
        """File kind is classified by extension."""
        assert resume_file_kind("cv.pdf") == "pdf"
        assert resume_file_kind("cv.docx") == "docx"
        assert resume_file_kind("cv.doc") == "doc"
