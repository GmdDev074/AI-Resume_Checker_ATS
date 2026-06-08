"""Microsoft Word (.docx) parsing service."""

from io import BytesIO
from pathlib import Path
from typing import BinaryIO, Tuple, Union

from resume_analyzer.utils.logger import get_logger
from resume_analyzer.utils.text_cleaner import clean_text

logger = get_logger(__name__)


class DocxParser:
    """Extract and validate text from resume Word (.docx) files."""

    def validate_docx(self, source: Union[bytes, BinaryIO, Path]) -> Tuple[bool, str]:
        """
        Validate that the input is a readable .docx file.

        Args:
            source: File bytes, stream, or path.

        Returns:
            Tuple of (is_valid, message).
        """
        try:
            doc = self._open_document(source)
            if not doc.paragraphs and not doc.tables:
                return False, "Word document appears to be empty."
            return True, "Valid Word document."
        except Exception as exc:
            logger.warning("DOCX validation failed: %s", exc)
            return False, f"Invalid or corrupted Word document: {exc}"

    def extract_text(self, source: Union[bytes, BinaryIO, Path]) -> str:
        """
        Extract and clean all text from a Word resume.

        Args:
            source: File bytes, stream, or path.

        Returns:
            Cleaned resume text.

        Raises:
            ValueError: If extraction fails.
        """
        try:
            doc = self._open_document(source)
            parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            raw = "\n".join(parts)
            if not raw.strip():
                raise ValueError("Word document contains no extractable text.")
            return clean_text(raw)
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"Failed to extract Word document text: {exc}") from exc

    def _open_document(self, source: Union[bytes, BinaryIO, Path]):
        """Open document from various input types."""
        from docx import Document

        if isinstance(source, Path):
            return Document(str(source))
        if isinstance(source, bytes):
            return Document(BytesIO(source))
        if hasattr(source, "read"):
            data = source.read()
            if hasattr(source, "seek"):
                source.seek(0)
            return Document(BytesIO(data))
        raise TypeError("Unsupported Word document source type.")

    def extract_from_upload(self, uploaded_bytes: bytes) -> str:
        """
        Convenience method for Streamlit uploaded file bytes.

        Args:
            uploaded_bytes: Raw file content.

        Returns:
            Extracted text.
        """
        valid, message = self.validate_docx(uploaded_bytes)
        if not valid:
            raise ValueError(message)
        return self.extract_text(BytesIO(uploaded_bytes))
